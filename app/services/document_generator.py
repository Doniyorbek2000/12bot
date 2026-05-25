import os
import uuid
import logging
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.config import settings

logger = logging.getLogger("adolat_ai_bot.document_generator")

class DocumentGenerator:
    def __init__(self, output_dir: str = "generated_docs"):
        self.output_dir = output_dir
        # Create output directory if it doesn't exist
        os.makedirs(self.output_dir, exist_ok=True)

    def create_docx(self, document_type: str, document_text: str) -> str:
        """Matn asosida professional yuridik DOCX hujjati yaratish"""
        try:
            doc = docx.Document()
            
            # Page margins setup
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(0.6)

            # Font setup
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Split Gemini text into lines and format
            lines = document_text.strip().split("\n")
            
            # Biz matndagi header, kimga/kimdan qismlarini chiroyli formatlashimiz mumkin
            header_section = True
            
            for line in lines:
                line_str = line.strip()
                if not line_str:
                    continue
                
                # Kimga, Kimdan qismlari (odatda o'ng tomonga moslashtiriladi)
                if header_section and any(keyword in line_str.lower() for keyword in ["kimga:", "kimdan:", "yuboruvchi:", "oluvchi:", "raisi:", "prokuroriga:", "sudiga:"]) or (header_section and line_str.startswith("-") and ":" in line_str):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run(line_str)
                    run.font.size = Pt(11)
                    run.italic = True
                
                # Hujjat sarlavhasi (O'rtaga moslashtirilgan, Bold)
                elif any(title in line_str.upper() for title in ["ARIZA", "SHIKOYAT", "DA'VO ARIZASI", "TUSHUNTIRISH XATI", "MUROJAAT", "TALABNOMA"]) and len(line_str) < 50:
                    header_section = False # Header tugadi, asosiy matn boshlandi
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(18)
                    p.paragraph_format.space_after = Pt(12)
                    run = p.add_run(line_str)
                    run.bold = True
                    run.font.size = Pt(14)
                
                # Asosiy matn
                else:
                    header_section = False
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(6)
                    
                    # Agar satr imzo yoki sana bo'lsa
                    if any(x in line_str.lower() for x in ["sana:", "imzo:", "yil", "2026", "__"]):
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        p.paragraph_format.space_before = Pt(18)
                    
                    run = p.add_run(line_str)

            # Generate unique filename
            filename = f"{document_type}_{uuid.uuid4().hex[:8]}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # Save file
            doc.save(filepath)
            logger.info(f"DOCX hujjat yaratildi va saqlandi: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"DOCX yaratishda xatolik: {e}")
            raise ValueError(f"Hujjat faylini yaratishda xatolik yuz berdi: {str(e)}")

    def create_pdf(self, docx_path: str) -> str:
        """
        Kelgusida DOCX faylini PDF ga o'tkazish uchun tayyor arxitektura.
        Windows yoki Linux deployiga qarab LibreOffice yoki python PDF kutubxonalari orqali amalga oshiriladi.
        """
        # Hozircha DOCX yo'lini o'zini yoki skeleton PDF yo'lini qaytaramiz
        pdf_path = docx_path.replace(".docx", ".pdf")
        # Real loyihada bu yerda convertion logikasi yoziladi
        return pdf_path
