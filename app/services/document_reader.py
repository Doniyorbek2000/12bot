import os
import logging
from io import BytesIO
import PyPDF2
import docx
from app.config import settings

logger = logging.getLogger("adolat_ai_bot.document_reader")

class DocumentReader:
    @staticmethod
    def read_txt(file_bytes: bytes) -> str:
        """TXT fayldan matnni o'qish, turli xil kodlashlarni qo'llab-quvvatlaydi"""
        encodings = ["utf-8", "utf-16", "latin-1", "cp1251"]
        for encoding in encodings:
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Matn faylini o'qib bo'lmadi. Kodlash formati to'g'ri kelmadi.")

    @staticmethod
    def read_pdf(file_bytes: bytes) -> str:
        """PDF fayldan matnni o'qish (PyPDF2 yordamida)"""
        try:
            pdf_file = BytesIO(file_bytes)
            reader = PyPDF2.PdfReader(pdf_file)
            extracted_text = ""
            
            # Cheklov: Juda katta pdf bo'lsa dastlabki 30 sahifani olamiz (Token overflowdan saqlanish uchun)
            max_pages = min(len(reader.pages), 30)
            
            for page_num in range(max_pages):
                page = reader.pages[page_num]
                text = page.extract_text()
                if text:
                    extracted_text += f"\n--- Sahifa {page_num + 1} ---\n" + text
            
            if not extracted_text.strip():
                raise ValueError("PDF fayldan hech qanday matn ajratib bo'lmadi (Ehtimol skaner qilingan rasm).")
                
            return extracted_text
        except Exception as e:
            logger.error(f"PDF o'qishda xatolik: {e}")
            raise ValueError(f"PDF faylni o'qishda xatolik yuz berdi: {str(e)}")

    @staticmethod
    def read_docx(file_bytes: bytes) -> str:
        """DOCX fayldan matnni o'qish (python-docx yordamida)"""
        try:
            docx_file = BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            extracted_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text.append(paragraph.text)
            
            # Jadvallardan ham matn ajratish
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        extracted_text.append(" | ".join(row_text))
            
            result = "\n".join(extracted_text)
            if not result.strip():
                raise ValueError("DOCX fayldan matn topilmadi.")
                
            return result
        except Exception as e:
            logger.error(f"DOCX o'qishda xatolik: {e}")
            raise ValueError(f"DOCX faylni o'qishda xatolik yuz berdi: {str(e)}")

    def read_document(self, file_bytes: bytes, file_name: str) -> str:
        """Fayl kengaytmasini aniqlab, mos o'quvchi xizmatga yo'naltirish"""
        ext = os.path.splitext(file_name)[1].lower()
        if ext == ".txt":
            return self.read_txt(file_bytes)
        elif ext == ".pdf":
            return self.read_pdf(file_bytes)
        elif ext in [".docx", ".doc"]:
            return self.read_docx(file_bytes)
        else:
            raise ValueError(f"Qo'llab-quvvatlanmaydigan fayl turi: {ext}. Faqat PDF, DOCX va TXT fayllarni yuboring.")
